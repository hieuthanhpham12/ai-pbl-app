import streamlit as st
import pandas as pd
import csv
from sentence_transformers import SentenceTransformer, util
import matplotlib.pyplot as plt
from docx import Document

# Load prompt library from CSV
def load_prompt_library(csv_path):
    data = []
    with open(csv_path, encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            data.append({
                "question": row["question"],
                "method": row["method"],
                "explanation": row["explanation"],
                "code": row["code"]
            })
    return data

# Export assignment suggestions to Word
def export_to_word(dataset_summary, suggestions):
    doc = Document()
    doc.add_heading("Statistical Project Assignment", level=1)
    doc.add_paragraph("The system analyzed the dataset and suggested the following project topics:")
    doc.add_heading("1. Dataset Summary", level=2)
    doc.add_paragraph(dataset_summary)
    doc.add_heading("2. Suggested Tasks", level=2)
    for s in suggestions:
        doc.add_paragraph(f"- {s}", style='List Bullet')
    file_path = "project_assignment.docx"
    doc.save(file_path)
    return file_path

# Streamlit UI
st.set_page_config(page_title="AI-PBL for Statistics", layout="wide")
st.title("📊 AI-Powered Support for Teaching Probability and Statistics")

menu = st.sidebar.radio("🎯 Select Mode", ["👩‍🏫 Instructor", "👨‍🎓 Student", "📋 Rubric Evaluation"])

model = SentenceTransformer("paraphrase-MiniLM-L6-v2")
prompt_library = load_prompt_library("ai_agent_prompt_library_full.csv")
questions = [q["question"] for q in prompt_library]
question_embeddings = model.encode(questions, convert_to_tensor=True)

# Instructor Interface
if menu == "👩‍🏫 Instructor":
    st.header("👩‍🏫 Instructor Interface")
    uploaded_file = st.file_uploader("📤 Upload CSV Dataset", type="csv")
    if uploaded_file:
        df = pd.read_csv(uploaded_file)
        st.success(f"✅ Dataset loaded: {df.shape[0]} rows × {df.shape[1]} columns.")
        st.dataframe(df.head())

        st.subheader("📊 Quick Data Visualization")
        st.bar_chart(df.select_dtypes(include='number'))

        cat_cols = df.select_dtypes(include='object').columns.tolist()
        num_cols = df.select_dtypes(include='number').columns.tolist()
        suggestions = []

        if len(num_cols) >= 2:
            suggestions.append("🔍 Explore correlation between numerical variables.")
        if len(cat_cols) >= 2:
            suggestions.append("📊 Test independence between categorical variables.")
        if cat_cols and num_cols:
            suggestions.append("📈 Compare group means (t-test or ANOVA).")

        st.subheader("📌 Suggested Project Topics")
        for s in suggestions:
            st.markdown(f"- {s}")

        if st.button("📝 Export to Word"):
            summary = f"Rows: {df.shape[0]}, Columns: {df.shape[1]}\nNumerical: {num_cols}\nCategorical: {cat_cols}"
            file_path = export_to_word(summary, suggestions)
            with open(file_path, "rb") as f:
                st.download_button("📥 Download Assignment File", data=f, file_name="project_assignment.docx")

# Student Interface
elif menu == "👨‍🎓 Student":
    st.header("🤖 Student Interface")
    user_input = st.text_input("📌 Ask a question about statistical analysis:")
    if user_input:
        query_vec = model.encode(user_input, convert_to_tensor=True)
        scores = util.cos_sim(query_vec, question_embeddings)[0]
        top_idx = scores.argmax().item()
        match = prompt_library[top_idx]
        st.success(f"📘 Suggested Method: {match['method']}")
        st.markdown(f"**🧠 Explanation:** {match['explanation']}")
        st.code(match['code'], language="python")

# Rubric Evaluation Interface
elif menu == "📋 Rubric Evaluation":
    st.header("📋 Evaluate Group Using Rubric")
    st.markdown("Select performance level for each criterion:")
    criteria = {
        "Data Cleaning & Preprocessing": [],
        "Statistical Analysis": [],
        "Visualization & Presentation": []
    }
    levels = [
        "0 - Not achieved", "1 - Minimal", "2 - Basic", "3 - Expected", "4 - Excellent"
    ]
    scores = {}
    for k in criteria.keys():
        scores[k] = st.selectbox(f"{k}:", levels, index=2)

    if st.button("✅ Calculate Total Score"):
        total = sum([int(s[0]) for s in scores.values()])
        st.markdown(f"**🔢 Total Score: {total}/12**")
